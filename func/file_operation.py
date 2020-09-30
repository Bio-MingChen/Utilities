# -*- coding=utf-8 -*-

#Functions in this script are all about file operations
#like open files by their file type or get columns by file title etc.

import gzip
import docx

#=======================================================
#TitleParser is used to get columns by file's title name
#========================================================
class TitleParser():
    """
    Reading in a file title and than getting indicated column element.
    Title names are all CaseIgnore
    """
    
    def __init__(self,title):
        self.title_list = [i.lower() for i in title.strip().split('\t')]

    def get_field(self,line_list,colname,check=True):
        """
        Reading in a list and returning the element with the 
        index which colname in title's list
        """
        if check:
            if len(self.title_list) != len(line_list):
                raise Exception("Title length differs with line!")
        try:
            idx = self.title_list.index(str(colname).lower())
        except:
            print('{colname} not in title！'.format(colname=colname))
            return None

        return line_list[idx]

    def have_title(self,colname):
        """
        Judging whether a colname is in title
        """
        if str(colname).lower() in self.title_list:
            return True
        return False
    
    def get_idx(self,colname):
        """
        return index of columns by their name in title list
        """
        idx = self.title_list.index(str(colname).lower())

        return idx

#=======================================
#gopen function is used to open file by their suffix
# open .gz file with gzip.open else open
#=======================================

def gopen(filename,mode):
    """
    open .gz file with gzip.open
    open plain text file with open
    """
    if filename.endswith('.gz'):
        return gzip.open(filename,mode)

    return open(filename,mode)
#========================================
#file exist and not empty judgement
# gzip file is not empty even thought it has no content
#========================================
def file_exist_and_not_empty(f):
    """
    gzip file is not empty even though it has no content!
    """ 
    if os.path.exists(f):
        if f.endswith('gz') and (os.path.getsize(f) < 100):
            with gopen(f,'r') as indata:
                if indata.read():
                    return True 
        else: 
            if os.path.getsize(f) != 0:
                return True
    return False
#========================================
#Class DocxApi help to operate Words easier 
#like add paragraph,picture or table to the
# place you want to insert etc.
#========================================
class DocxApi():
    """
    docx api implement adding,replacing and deletion operation
    """
    def __init__(self,docx_file):
        self.doc = docx.Document(docx_file) if docx_file else docx.Document()

    def set_font(self,en_style='Times New Roman',zh_style='微软雅黑'):
        """
        set global font 
        """
        self.doc.styles['Normal'].font.name = en_style
        self.doc.styles['Normal']._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'),zh_style.decode('utf-8'))

    def set_para_font_size(self,paragraph,size=10.5):
        """
        set paragraph font size
        """
        for run in paragraph.runs:
            run.font.size = docx.shared.Pt(size)
        
        return paragraph

    def set_tab_font_size(self,table,size=10.5):
        """
        set table font size
        """
        for row in table.rows:
            for cell in row.cells:
                for pargraph in cell.paragraphs:
                    self.set_para_font_size(pargraph,size=size)
        return table

    def add_picture(self,pic_path,width=5.0):
        """
        In order to move picture to some place you want put them
        here this function will build a new paragprah and then 
        use the 'run.add_picture()' method to add picture. After that,
        you can move it like moving a paragraph to insert to docx or
        justify styles.
        
        Width 5.0 is the proper picture width in Test,you can change it
        to match your need.
        """
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(pic_path,width=docx.shared.Inches(width))
        return paragraph

    def add_page_break(self):
        """
        start from new page
        """
        self.doc.add_page_break()

    def delete_tab_para(self,table):
        """
        delete table or paragraph
        """
        table._element.getparent().remove(table._element)
        print('remove complete!')

    def add_paras(self,content):
        """
        add a new paragraph
        """
        return self.doc.add_paragraph(content.decode('utf-8'))

    def add_table(self,table_list):
        """
        add a new table
        table_list : [[1,2,3],[4,5,6],...]
        """
        nrow = len(table_list)
        ncol = len(table_list[0])
        print(nrow,ncol)
        table = self.doc.add_table(0,ncol)
        for row in table_list:
            row_cells = table.add_row().cells
            for idx,cell in enumerate(row):
                try:
                    row_cells[idx].text = cell.decode('utf-8')
                except:
                    row_cells[idx].text = cell
                    # print("{cell} of {row} can\'t be decode correctly!".format(cell=cell,row=row))
                    # print(cell,row)
                    # exit(1)
        return table

    def move_obj(self,obj,tag):
        """
        move the obj to the next of tag
        """
        if isinstance(obj,docx.table.Table):
            obj = obj._tbl
        elif isinstance(obj,docx.text.paragraph.Paragraph):
            obj = obj._p
        else:
            print('UnKnown type of input object!')
            exit(1)
        for paragraph in self.doc.paragraphs:
            if paragraph.text.startswith(tag.decode('utf-8')):
                paragraph._p.addnext(obj)
                print('add complete!')
    
    def save(self,ofile=None):
        if not ofile:
            ofile = 'test.docx'
        self.doc.save(ofile)

    def replace_para(self,paragraph,new_content):
        """
        clear paragraph and add new content
        """
        paragraph.clear()
        paragraph.add_run(new_content)
    
        return paragraph
    
    def get_pt_by_size(self,size):
        """
        transformation of pt and chinese size 号
        """
        pt_hao_trans_dict = {
                            '八号':5.0,
                            '七号':5.5,
                            '小六':6.5,
                            '六号':7.5,
                            '小五':9.0,
                            '五号':10.5,
                            '小四':12.0,
                            '四号':14.0,
                            '小三':13.0,
                            '三号':16.0,
                            '小二':18.0,
                            '二号':22.0,
                            '小一':24.0,
                            '一号':26.0,
                            '小初':36.0,
                            '初号':42.0,
                            }
        if size in pt_hao_trans_dict:
            # return docx.shared.Pt(pt_hao_trans_dict[size])
            return pt_hao_trans_dict[size]
        else:
            print('Unknown input size,size should be in {keys}'.format(keys=list(pt_hao_trans_dict)))
            return None
